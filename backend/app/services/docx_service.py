import hashlib
import io

from docx import Document


DOCX_MAX_SIZE = 50 * 1024 * 1024


class DocxResult:
    def __init__(self, title: str, text: str, index_key: str):
        self.title = title
        self.text = text
        self.index_key = index_key


def content_to_index_key(content: bytes) -> str:
    return "docx_" + hashlib.sha256(content).hexdigest()[:16]


def process_docx(file_bytes: bytes, title: str = "") -> DocxResult:
    if len(file_bytes) > DOCX_MAX_SIZE:
        raise ValueError(f"DOCX exceeds size limit of {DOCX_MAX_SIZE // (1024 * 1024)}MB.")

    doc = Document(io.BytesIO(file_bytes))

    paragraphs: list[str] = []
    first_heading = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading") and not first_heading:
            first_heading = text[:200]
        paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    full_text = "\n\n".join(paragraphs)
    if len(full_text) < 50:
        raise ValueError("Could not extract meaningful content from this DOCX file.")

    title = title.strip() or first_heading or "Untitled Document"
    index_key = content_to_index_key(file_bytes)

    return DocxResult(title=title, text=full_text, index_key=index_key)

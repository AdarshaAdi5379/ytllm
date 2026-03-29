import io
import os
import httpx
from datetime import datetime
from typing import Any
from docx import Document
from docx.shared import (
    Inches,
    Pt,
    RGBColor,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

from app.models import Message


class ExportData:
    def __init__(
        self,
        video_id: str,
        title: str,
        channel_name: str,
        duration: str,
        thumbnail_url: str,
        summary: str,
        chat_history: list[Message],
        include_transcript: bool = False,
        transcript: str | None = None,
    ):
        self.video_id = video_id
        self.title = title
        self.channel_name = channel_name
        self.duration = duration
        self.thumbnail_url = thumbnail_url
        self.summary = summary
        self.chat_history = chat_history
        self.include_transcript = include_transcript
        self.transcript = transcript


async def fetch_thumbnail(thumbnail_url: str) -> bytes | None:
    """Fetches a YouTube thumbnail and returns it as bytes."""
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(thumbnail_url, timeout=8.0)
            response.raise_for_status()
            return response.content
    except Exception:
        # Try fallback URL
        try:
            video_id_match = thumbnail_url.rsplit("/vi/", 1)[-1].split("/")[0]
            if video_id_match:
                fallback_url = (
                    f"https://img.youtube.com/vi/{video_id_match}/hqdefault.jpg"
                )
                async with httpx.AsyncClient() as client:
                    response = await client.get(fallback_url, timeout=5.0)
                    response.raise_for_status()
                    return response.content
        except Exception:
            pass
    return None


def format_timestamp(iso_timestamp: str) -> str:
    """Formats ISO timestamp to readable string."""
    try:
        return datetime.fromisoformat(iso_timestamp.replace("Z", "+00:00")).strftime(
            "%Y-%m-%d %H:%M:%S"
        )
    except Exception:
        return iso_timestamp


async def generate_pdf(data: ExportData) -> bytes:
    """Generates a PDF export and returns it as bytes."""
    pdf = PDF()
    pdf.add_page()

    # Colors
    USER_BG = (227, 242, 253)  # #E3F2FD
    AI_BG = (245, 245, 245)  # #F5F5F5
    ACCENT = (21, 101, 192)  # #1565C0
    HEADER_BG = (26, 35, 126)  # #1A237E

    # Cover section - header bar
    pdf.set_fill_color(*HEADER_BG)
    pdf.rect(0, 0, 210, 40, "F")

    pdf.set_text_color(255, 255, 255)
    pdf.set_font("helvetica", "B", 18)
    pdf.cell(0, 25, "YouTube AI Chat Export", ln=True, align="C")

    pdf.set_font("helvetica", "", 10)
    pdf.ln(5)
    pdf.cell(
        0, 10, f"Exported on {datetime.now().strftime('%Y-%m-%d')}", ln=True, align="C"
    )

    pdf.ln(15)

    # Thumbnail and metadata
    thumbnail_bytes = await fetch_thumbnail(data.thumbnail_url)
    x_start = 10
    y_start = pdf.get_y()

    if thumbnail_bytes:
        tmp_path = None
        try:
            # Save to temp file for FPDF
            import tempfile

            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                tmp.write(thumbnail_bytes)
                tmp_path = tmp.name

            pdf.image(tmp_path, x=x_start, y=y_start, w=60, h=34)
            meta_x = 75
        except Exception:
            meta_x = x_start
        finally:
            if tmp_path:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
    else:
        meta_x = x_start

    # Write metadata aligned next to thumbnail (FPDF2 doesn't support `x=` on cell)
    pdf.set_xy(meta_x, y_start)

    pdf.set_text_color(0, 0, 0)
    pdf.set_font("helvetica", "B", 14)
    pdf.set_text_color(*ACCENT)
    pdf.cell(0, 10, data.title, ln=True)

    pdf.set_font("helvetica", "", 10)
    pdf.set_text_color(85, 85, 85)
    pdf.set_x(meta_x)
    pdf.cell(0, 8, f"Channel: {data.channel_name}", ln=True)
    pdf.set_x(meta_x)
    pdf.cell(0, 8, f"Duration: {data.duration}", ln=True)
    pdf.set_x(meta_x)
    pdf.cell(0, 8, f"Video ID: {data.video_id}", ln=True)

    pdf.ln(15)

    # Summary section
    pdf.set_draw_color(204, 204, 204)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    pdf.set_font("helvetica", "B", 12)
    pdf.set_text_color(*ACCENT)
    pdf.cell(0, 10, "Video Summary", ln=True)

    pdf.set_font("helvetica", "", 9)
    pdf.set_text_color(51, 51, 51)
    pdf.multi_cell(0, 6, data.summary)

    pdf.ln(10)

    # Conversation section
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    pdf.set_font("helvetica", "B", 12)
    pdf.set_text_color(*ACCENT)
    pdf.cell(0, 10, "Conversation", ln=True)
    pdf.ln(5)

    for msg in data.chat_history:
        is_user = msg.role == "user"
        bg_color = USER_BG if is_user else AI_BG
        label = "You" if is_user else "AI Assistant"
        label_color = ACCENT if is_user else (46, 125, 50)  # #2E7D32

        # Check for new page
        if pdf.get_y() > 240:
            pdf.add_page()

        # Message bubble background
        pdf.set_fill_color(*bg_color)
        pdf.rect(10, pdf.get_y(), 190, 25, "F")

        pdf.set_font("helvetica", "B", 8)
        pdf.set_text_color(*label_color)
        pdf.cell(15, 8, label, ln=False)

        pdf.set_font("helvetica", "", 7)
        pdf.set_text_color(136, 136, 136)
        pdf.cell(170, 8, format_timestamp(msg.timestamp), ln=True, align="R")

        pdf.set_font("helvetica", "", 9)
        pdf.set_text_color(34, 34, 34)
        pdf.multi_cell(185, 5, msg.content)

        pdf.ln(8)

    # Optional transcript
    if data.include_transcript and data.transcript:
        pdf.add_page()
        pdf.set_font("helvetica", "B", 14)
        pdf.set_text_color(*ACCENT)
        pdf.cell(0, 10, "Full Transcript", ln=True)
        pdf.ln(5)

        pdf.set_font("helvetica", "", 8)
        pdf.set_text_color(51, 51, 51)
        pdf.multi_cell(0, 4, data.transcript)

    # Output
    pdf_out = pdf.output(dest="S")
    if isinstance(pdf_out, (bytes, bytearray)):
        return bytes(pdf_out)
    return str(pdf_out).encode("latin-1")


async def generate_docx(data: ExportData) -> bytes:
    """Generates a DOCX export and returns it as bytes."""
    doc = Document()

    # Title
    title_para = doc.add_heading(data.title, level=1)

    # Metadata
    doc.add_paragraph(f"Channel: {data.channel_name}")
    doc.add_paragraph(f"Duration: {data.duration}")
    doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()

    # Summary
    doc.add_heading("Video Summary", level=2)
    doc.add_paragraph(data.summary)
    doc.add_paragraph()

    # Conversation
    doc.add_heading("Conversation", level=2)

    for msg in data.chat_history:
        is_user = msg.role == "user"
        label = "You" if is_user else "AI Assistant"

        if is_user:
            # User messages with shading
            table = doc.add_table(rows=3, cols=1)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            cell = table.rows[0].cells[0]
            shading_elem = OxmlElement("w:shd")
            shading_elem.set(qn("w:fill"), "E3F2FD")
            cell._tc.get_or_add_tcPr().append(shading_elem)

            cell.text = f"{label}\n{msg.content}\n{format_timestamp(msg.timestamp)}"
        else:
            # AI messages as paragraphs
            p = doc.add_paragraph()
            p.add_run(f"{label}\n").bold = True
            p.add_run(msg.content)
            p.add_run(f"\n{format_timestamp(msg.timestamp)}").font.color.rgb = RGBColor(
                136, 136, 136
            )
            doc.add_paragraph()

    # Transcript
    if data.include_transcript and data.transcript:
        doc.add_page_break()
        doc.add_heading("Full Transcript", level=2)
        doc.add_paragraph(data.transcript)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


class PDF(FPDF):
    def header(self):
        pass

    def footer(self):
        self.set_y(-15)
        self.set_font("helvetica", "", 8)
        self.set_text_color(136, 136, 136)
        self.cell(0, 10, "YouTube AI Chat Export", align="C")
